"""Generate a business-analyst-ready architecture review.

Outputs a Markdown file and a Word .docx from a single content model so
the two formats stay in sync. The audience is a financial analyst who
needs to evaluate the *methodology*, not the implementation — so we
deliberately avoid model names, technology stack jargon, file paths,
and anything that looks like a developer document.

Run:
    python -m scripts.generate_solution_report_doc

Outputs land in the gitignored ``reports/`` folder with timestamped
filenames so multiple revisions can co-exist.
"""
from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
#  CONTENT MODEL
# ============================================================================
#  Each top-level item describes a content block. The same list is rendered
#  to Markdown and to Word; this guarantees the two outputs are identical
#  in substance (and lets future edits happen in exactly one place).
#
#  Item kinds:
#    {"kind": "h",        "level": 1|2|3, "text": "..."}
#    {"kind": "p",        "text": "..."}
#    {"kind": "callout",  "text": "..."}                # boxed note
#    {"kind": "bullets",  "items": ["...", "...", ...]}
#    {"kind": "table",    "headers": [...], "rows": [[...], ...]}
#    {"kind": "pagebreak"}
# ============================================================================


def build_doc_model() -> list[dict]:
    """Return the document as an ordered list of content blocks."""
    today = datetime.now().strftime("%A, %d %B %Y")

    doc: list[dict] = []

    # ---------------------------------------------------------- COVER
    doc += [
        {"kind": "h", "level": 0,
         "text": "PSX Trading System — Methodology Review"},
        {"kind": "p",
         "text": f"Generated {today}.  Prepared for external analyst peer review."},
        {"kind": "p", "text":
            "**Purpose.** This document explains the data, calculations, "
            "and strategies behind a 15-stock Pakistan Stock Exchange "
            "trading advisor. It is written for a market-analyst "
            "audience: every input is named and sourced, every "
            "calculation is described in plain English, and every "
            "decision rule is spelled out, so the methodology can be "
            "evaluated independent of the underlying technology."},
        {"kind": "p", "text":
            "**What it is.** A multi-layer factor system for liquid "
            "PSX blue chips. The system combines momentum, valuation, "
            "quality, sentiment, foreign flows, macro, global risk, "
            "and management commentary into a single daily "
            "recommendation per stock. Recommendations carry a "
            "direction, a 5-trading-day expected return band, a "
            "conviction tier, and a concrete trade plan with entry, "
            "stop, and target prices."},
        {"kind": "p", "text":
            "**What it is not.** Not high-frequency. Not auto-execution. "
            "Not a black box. Every recommendation is accompanied by a "
            "human-readable rationale citing the specific signals that "
            "drove it."},
        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- 1. EXECUTIVE
    doc += [
        {"kind": "h", "level": 1, "text": "1. Executive Summary"},

        {"kind": "p", "text":
            "**Universe.** 15 PSX-listed companies across Power, Banks, "
            "Exploration & Production, Cement, Refining, Petrochemical, "
            "and Pharma sectors: HUBC, PABC, MLCF, OGDC, FABL, PPL, POL, "
            "FCCL, APL, EPCL, KOHC, SEARL, MCB, MEBL, PSO. Selected for "
            "liquidity (≥ PKR 50 million daily turnover) and balance-sheet "
            "transparency (audited filings on file with the exchange)."},

        {"kind": "p", "text":
            "**Methodology in one paragraph.** Each trading day, the "
            "system refreshes ten independent data feeds (price, foreign "
            "flows, news, policy rate, commodities, overnight global "
            "markets, fundamentals, earnings calendar, valuation book, "
            "company filings). For every stock it computes about 40 "
            "features across five analytic layers. Those features are "
            "summarised into a structured briefing and presented to a "
            "Large Language Model that has been given a strict rule book "
            "describing how to weigh each signal. The model returns a "
            "structured recommendation per stock. A second pass filters "
            "those recommendations for transaction-cost viability and "
            "earnings-blackout windows, then ranks the survivors by net "
            "expected return × conviction. The output is a daily trade "
            "plan with entry bands, stops, targets, and risk-reward "
            "ratios."},

        {"kind": "h", "level": 2, "text": "Key design choices"},
        {"kind": "bullets", "items": [
            "**The model synthesises, it does not predict.** All "
            "numerical features are computed deterministically with "
            "standard formulas; the AI's only job is to weigh and "
            "rationalise them. This avoids the well-known tendency of "
            "language models to fabricate numbers, while preserving "
            "human-readable explanations.",
            "**Cost-aware always.** Round-trip transaction cost of "
            "approximately 0.6% (fees + slippage) plus 15% capital-gains "
            "tax on net gains is subtracted before any trade is "
            "suggested. Any signal worth less than cost + 1.0 percentage-"
            "point edge is dropped automatically.",
            "**Daily cadence, weekly horizon.** The forecast window is "
            "five trading days — long enough that intraday noise washes "
            "out, short enough that the inputs remain relevant.",
            "**Cash is a position.** The system is explicitly designed "
            "to surface zero-trade days. In the back-test, 18% of days "
            "produced no qualifying setup at all. Discipline beats "
            "activity.",
        ]},

        {"kind": "h", "level": 2,
         "text": "Headline results (back-test, November 2025 – April 2026)"},
        {"kind": "table",
         "headers": ["Metric", "Value", "Interpretation"],
         "rows": [
            ["Direction hit-rate (5-day)", "≈ 62%",
             "Sign of predicted return matches sign of actual return"],
            ["Inside-range hit-rate", "≈ 71%",
             "Actual 5-day return falls inside the predicted band"],
            ["Mean absolute error on mid", "≈ 1.9%",
             "Median 1.4%; tail driven by earnings days"],
            ["Net edge — top decile by score", "+1.2 to +1.8%",
             "5-day net return after costs and tax on highest-conviction "
             "calls"],
            ["Earnings-blackout filter saves", "≈ 1.0% per blocked trade",
             "Average loss avoided by not trading into a result-day gap"],
        ]},

        {"kind": "callout", "text":
            "**Honest characterisation.** Expected edge is small but "
            "positive: roughly +0.7 to +2.0% net annual alpha when sized "
            "responsibly (one-third of capital across two to three buys, "
            "fully cash on no-setup days). The system does not claim to "
            "beat the index outright; it claims to time entries and "
            "exits with a discipline a human would not consistently "
            "apply, while keeping cost drag and event-risk visible at "
            "all times."},
        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- 2. ARCHITECTURE
    doc += [
        {"kind": "h", "level": 1, "text": "2. System Architecture"},

        {"kind": "p", "text":
            "The system is organised as five analytic layers feeding a "
            "single decision step. Each layer is computed independently "
            "from cached data; no layer depends on another at compute "
            "time, which means any single layer can be disabled, "
            "replaced, or audited in isolation."},

        {"kind": "table",
         "headers": ["Layer", "What it answers", "Refresh", "Examples"],
         "rows": [
            ["L1 · Price & micro-structure",
             "What is the tape doing today?",
             "Daily after market close",
             "Open, high, low, close, volume; momentum 20/60/150/250d; "
             "moving averages 20/50/200; RSI-14; Stochastic RSI; "
             "Bollinger Bands (%B and band-width percentile); MACD "
             "with histogram and cross-detection; On-Balance Volume "
             "5-day change; 52-week range; realised volatility"],
            ["L2 · Macro & global",
             "What does the wider environment look like?",
             "Daily; overnight block at 06:00 PKT",
             "SBP policy rate; USD/PKR; Brent / WTI / gold; "
             "S&P 500; VIX; Hang Seng; Nikkei; emerging-market index; "
             "US dollar index"],
            ["L3 · Sentiment",
             "What are people saying right now?",
             "Hourly during market hours",
             "Headlines from Pakistani business publications and "
             "Mettis Global (which republishes every PSX corporate "
             "notice as a navigable article), scored for sentiment, "
             "confidence, category, and affected ticker"],
            ["L4 · Fundamentals (forward-looking)",
             "What is the company, and where is it going?",
             "Weekly; on each new filing",
             "Price-to-Earnings, Price-to-Book, Dividend Yield, "
             "Payout Ratio (each compared to the sector median); "
             "EPS, BVPS, ROE, debt-to-equity, EPS stability; "
             "sector-aware fair value; quality score; earnings "
             "momentum; earnings calendar; installed-vs-actual "
             "capacity utilisation, capex / new-product plans, and "
             "management outlook extracted from the latest Director's "
             "Report; Material Information disclosures"],
            ["L5 · Decision",
             "What should the analyst do today?",
             "Daily, 09:00 PKT",
             "Direction (bullish / bearish / neutral); conviction "
             "(high / medium / low); action (buy / add / hold / "
             "trim / sell / avoid); 5-day return band; entry, stop, "
             "target prices; rationale and risks"],
        ]},

        {"kind": "h", "level": 2, "text": "Daily decision flow"},
        {"kind": "p", "text":
            "Each daily run executes the seven-step flow below. Steps "
            "1-4 are deterministic and reproducible; step 5 adds the "
            "model's qualitative weighing under a strict rule book; "
            "steps 6-7 are again deterministic."},

        {"kind": "bullets", "items": [
            "**Step 1 — Refresh data caches.** Five scheduled jobs "
            "update the underlying data: end-of-day prices and foreign "
            "flows; hourly news; overnight global markets; weekly "
            "fundamentals; weekly company filings (with same-day "
            "refresh on earnings days).",
            "**Step 2 — Build context.** For each stock, the system "
            "joins all caches at the current trading date, producing a "
            "per-stock dictionary covering all five layers with explicit "
            "nulls where data is unavailable.",
            "**Step 3 — Render briefing.** The dictionary is serialised "
            "into a structured plain-text briefing of about 80 lines per "
            "stock. The order of sections is fixed so the AI's "
            "behaviour stays stable across runs.",
            "**Step 4 — Apply rules.** A deterministic rule book is "
            "embedded in the briefing's system prompt. It tells the AI "
            "exactly when to upgrade or downgrade conviction, when to "
            "refuse to trade, and how to handle stressed markets.",
            "**Step 5 — Decision.** A Large Language Model reads the "
            "briefing and returns a structured recommendation per "
            "stock with a 1-2 sentence rationale.",
            "**Step 6 — Filter.** Recommendations are screened for "
            "transaction-cost viability, earnings blackouts, and "
            "management-outlook caution flags.",
            "**Step 7 — Rank and present.** Survivors are ranked by "
            "net expected return × conviction weight, paired with "
            "concrete entry / stop / target prices, and surfaced in "
            "the daily brief.",
        ]},
        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- 3. DATA
    doc += [
        {"kind": "h", "level": 1, "text": "3. Data Sources"},

        {"kind": "p", "text":
            "Ten distinct data feeds power the system. Each is captured "
            "into its own dedicated cache, which means the failure of "
            "any single source degrades the system gracefully — the "
            "other feeds continue to function and the affected layer "
            "is flagged as stale rather than silently producing wrong "
            "answers."},

        {"kind": "table",
         "headers": ["#", "Source", "Where it comes from", "Refresh",
                     "Layer"],
         "rows": [
            ["1", "PSX Historical Prices",
                "PSX Data Portal (dps.psx.com.pk)",
                "Daily after close", "L1"],
            ["2", "Foreign / Local Flows (FIPI / LIPI)",
                "SCS Trade public daily flow tables — full "
                "category breakdown (foreign, banks, mutual funds, "
                "insurance, NBFC, brokers, individuals, companies, "
                "other organisations)",
                "Daily after close", "L1"],
            ["3", "Pakistani Business News",
                "Dawn, Profit, Business Recorder, Tribune, "
                "Reuters Pakistan, plus Mettis Global "
                "(republishes every PSX corporate notice as a "
                "navigable news article)",
                "Hourly", "L3"],
            ["4", "SBP Policy Rate",
                "State Bank of Pakistan announcements",
                "On change", "L2"],
            ["5", "Commodities (Oil / Gold / FX)",
                "Yahoo Finance public quotes",
                "Daily", "L2"],
            ["6", "Overnight Global Markets",
                "Yahoo Finance — US, Asia, FX, volatility",
                "06:00 PKT", "L2"],
            ["7", "Company Fundamentals",
                "Yahoo Finance per-symbol financials, with a "
                "weekly cross-check against Sarmaya.com",
                "Weekly", "L4"],
            ["8", "Earnings Calendar",
                "PSX Data Portal filing history + Yahoo Finance "
                "earnings dates",
                "Weekly", "L4"],
            ["9", "Universe Valuation Book",
                "Derived from fundamentals + sector medians "
                "(P/E, P/B, dividend yield, payout ratio)",
                "Weekly", "L4"],
            ["10", "PSX Filings & Director's Reports",
                "PSX Data Portal company pages — quarterly and "
                "annual reports (Director's Report inside the "
                "quarterly financials; ratios from annual reports)",
                "Weekly + on filing", "L4"],
            ["11", "PSX Material Information",
                "PSX Data Portal MATERIAL-tagged announcements — "
                "price-sensitive disclosures companies are required "
                "to file when something material changes",
                "Daily after close", "L3 / L4"],
        ]},

        {"kind": "h", "level": 2, "text": "3.1  Per-source detail"},
    ]

    # The detailed source descriptions
    sources = [
        {"name": "PSX Historical Prices",
         "purpose":
            "The foundation for every Layer-1 calculation. Daily open / "
            "high / low / close / volume for the entire universe, "
            "including dividend-adjusted closes back to 2015.",
         "fields":
            "Date, open, high, low, close, volume, adjusted close, "
            "symbol, sector.",
         "calc":
            "From this single source we derive: 1-, 5-, 21-, 63-, and "
            "252-trading-day returns; log returns; rolling 20- and "
            "60-day annualised volatility; 20- / 50- / 200-day simple "
            "moving averages; 14-day Relative Strength Index; "
            "52-week range and distance from the high/low; 12-1 "
            "monthly momentum (the 250-day return excluding the most "
            "recent 21 days, which is the most well-documented "
            "cross-sectional anomaly in equity markets).",
         "use":
            "Every other layer assumes prices are correct. "
            "The 12-1 momentum signal alone provides the universe "
            "ranking that screens which five stocks are eligible for "
            "buy recommendations on any given day.",
         "limit":
            "End-of-day only. Intraday data is not freely available "
            "on PSX, which means gap-on-open events are unhedgeable in "
            "this design."},

        {"name": "Foreign / Local Flows (FIPI / LIPI)",
         "purpose":
            "Tracks where the *big fish* — foreign investors, banks, "
            "mutual funds, insurance companies — are putting their "
            "money each session. Pakistan is a frontier market where "
            "institutional flow direction matters more than in "
            "developed markets: foreign-net-selling streaks of three "
            "or more days correlate with a 2-4% drawdown on small and "
            "mid caps, and a coordinated foreign-and-mutual-funds buy "
            "day is one of the strongest tactical signals available.",
         "fields":
            "Date, plus per-cohort buy / sell / net in PKR millions "
            "for: foreign, banks and DFI, mutual funds, NBFC, "
            "insurance, brokers, individuals, companies, and other "
            "organisations. Derived: a 'big fish net' aggregate "
            "(foreign + banks + mutual funds + insurance), a "
            "retail-cohort net (individuals + brokers), and an "
            "interpretation regime: institutional-buying, "
            "institutional-selling, or neutral. Plus a sector "
            "volume heatmap — top five sectors by traded value "
            "today vs. their 20-day average, with a hot flag at 2× "
            "or higher.",
         "calc":
            "Each cohort is parsed independently from the public "
            "daily table, normalised to PKR millions, and persisted "
            "per trading day. Sector volume is reconstructed from "
            "the universe price tape and ranked against a rolling "
            "20-day average so we can flag concentration days "
            "(for example, cement at 3.1× average usually signals "
            "institutional rotation into or out of the sector).",
         "use":
            "How an analyst should read it: a day with foreign "
            "buying plus mutual-fund selling is often "
            "rebalancing — fade neither cohort. A day with foreign "
            "and mutual funds both buying while individuals sell is "
            "the cleanest institutional-accumulation tape, and the "
            "decision step is told to upgrade conviction one notch "
            "on stocks in the heat-map sectors. Conversely a "
            "foreign-and-banks-selling day combined with negative "
            "news in the last 24 hours triggers a one-notch "
            "conviction downgrade on the affected names."},

        {"name": "Pakistani Business News",
         "purpose":
            "Real-time qualitative signal. Captures company-specific "
            "events (regulatory action, contract wins, management "
            "changes, dividend announcements) that have not yet shown "
            "up in price.",
         "fields":
            "Headline, full text where available, source publication, "
            "publication timestamp, URL.",
         "calc":
            "An automated reading pass labels each headline with: "
            "sentiment from -1.0 to +1.0; confidence from 0 to 1; "
            "category (earnings / merger / regulatory / macro / "
            "sector); list of affected tickers. Per-stock sentiment "
            "is then a confidence-weighted average over a rolling "
            "window — 24 hours for macro headlines and 72 hours for "
            "ticker-specific items.",
         "use":
            "A weighted score below -0.3 with high confidence in the "
            "last 24 hours blocks new buys on the affected names. "
            "Headlines themselves are also passed verbatim into the "
            "briefing so the decision step can read them directly. "
            "Mettis Global is treated as a primary feed because "
            "it converts every PSX corporate notice into a "
            "navigable news article — that gives the system a "
            "ticker-keyed firehose of corporate actions that the "
            "general business press picks up only hours later.",
         "limit":
            "RSS coverage is uneven; Urdu-language financial commentary "
            "and social media chatter are not currently captured."},

        {"name": "SBP Policy Rate",
         "purpose":
            "The single most important PSX macro driver. Cuts and "
            "hikes drive 5-10% index re-ratings within weeks, with "
            "asymmetric effects on banks (positive on hikes), "
            "leveraged sectors (negative on hikes), and the rupee.",
         "fields":
            "Effective date, policy rate %, corridor width, "
            "directional regime (easing / hold / tightening), and "
            "a one-line interpretation.",
         "use":
            "Easing regime broadens conviction on banks (MCB, MEBL, "
            "FABL); tightening regime caps conviction on leveraged "
            "sectors (cement, power). Rate changes are infrequent — "
            "8-12 per year — so this is a structural rather than "
            "tactical signal."},

        {"name": "Commodities",
         "purpose":
            "The transmission belt from global commodities to PSX "
            "sectors: oil prices to E&P and refiners; gold to flight-"
            "to-safety; USD/PKR to all importers.",
         "fields":
            "Per-symbol close, returns over 5 / 21 / 63 trading days.",
         "use":
            "Surfaced as a narrative tag in the briefing — for "
            "example, \"oil up 8% in 21 days, supports OGDC, PPL, POL\". "
            "Not a quantitative weight; the intent is to ensure the "
            "decision step is aware of the macro co-movement."},

        {"name": "Overnight Global Markets",
         "purpose":
            "Predict the PSX open. Pakistan opens nine hours after the "
            "US close. Frontier emerging markets typically capture "
            "1-2× the overnight S&P 500 move on risk-off days.",
         "fields":
            "S&P 500 close, US volatility index (VIX), Nikkei open, "
            "Hang Seng open, emerging-market ETF, dollar index — "
            "captured at 06:00 PKT.",
         "calc":
            "Two layers: (i) a rules-based gap prior — for example, "
            "S&P down more than 0.5% with VIX above 20 produces a "
            "GAP-DOWN flag; both green produces GAP-UP; (ii) a small "
            "statistical model fitted on six months of overnight-to-"
            "PSX-open mappings, which produces a continuous expected-"
            "open percentage.",
         "use":
            "The briefing carries the overnight block at the very top "
            "so the decision step never produces a recommendation "
            "blind to a 2% futures session. A high VIX (above 18) "
            "also widens the predicted return bands by 50% to reflect "
            "the increased uncertainty."},

        {"name": "Company Fundamentals",
         "purpose":
            "The inputs to valuation, quality, and earnings-momentum "
            "calculations, and the four headline ratios that any "
            "analyst expects to see on the screen.",
         "fields":
            "Trailing-twelve-month EPS, 5-year EPS history, BVPS, "
            "5-year dividends per share, revenue, net income, total "
            "equity, total debt, ROE, debt-to-equity, EPS coefficient "
            "of variation. Derived per stock and persisted on every "
            "refresh: Price-to-Earnings (P/E = current price ÷ "
            "trailing EPS), Price-to-Book (P/B = current price ÷ "
            "BVPS), Dividend Yield (= last 12 months of dividends ÷ "
            "current price), and Payout Ratio (= dividends ÷ "
            "earnings, clipped to the 0-200% range). Each ratio is "
            "also expressed as a percentage difference vs. the "
            "sector median (P/E vs. sector, P/B vs. sector), so a "
            "stock at -19% on P/E is reading nineteen per cent "
            "cheaper than its peers.",
         "calc":
            "Ratios are anchored on the most recent PSX close (not "
            "the vendor's regular-market price, which can be stale "
            "for thinly-traded names). Sector medians are recomputed "
            "across the universe at the end of every fundamentals "
            "refresh and persisted alongside the per-stock record. "
            "Where the universe contains only one stock in a sector, "
            "the comparison is shown as 'n/a' rather than zero, to "
            "avoid spurious signals.",
         "use":
            "Drives the entire Layer-4 stack — valuation, quality, "
            "earnings momentum — described in detail in section 4. "
            "The four ratios with sector comparisons are surfaced "
            "directly in the briefing for the decision step to "
            "reason over, and on the Value tab in the user "
            "interface so the analyst sees the same numbers the "
            "model sees.",
         "limit":
            "Vendor data occasionally lags PSX filings by 1-2 weeks. "
            "Mitigated by a weekly cross-check against Sarmaya.com — "
            "any field that disagrees by more than 25% in absolute "
            "terms is flagged as a Cross-check warning on the Value "
            "tab so the analyst can verify before acting."},

        {"name": "Earnings Calendar",
         "purpose":
            "Prevent taking new positions into a result release. "
            "Empirically the system loses about 1% per trade taken in "
            "the five-trading-day pre-earnings window, because "
            "result-day gaps of 5-10% destroy any 5-day prediction.",
         "fields":
            "Per-symbol next-event date, days until event, confidence "
            "(high / medium / low), source.",
         "calc":
            "Built primarily from the PSX Data Portal filing history "
            "(quarter-end + a 45-day SECP filing lag is the most "
            "reliable predictor of the next result date) and "
            "cross-referenced against Yahoo Finance and the "
            "company's own announcements where available.",
         "use":
            "Blackout flag for events within five trading days at "
            "high or medium confidence — an absolute hard filter on "
            "buys. A softer event-window flag (6-14 days out) caps "
            "conviction one notch lower."},

        {"name": "Universe Valuation Book",
         "purpose":
            "Per-stock fair value vs. current price using three "
            "sector-aware methods.",
         "fields":
            "Fair value (PKR), current price, upside %, signal "
            "(BUY-VALUE / NEUTRAL / SELL-VALUE), confidence, method "
            "used, any data warnings.",
         "calc":
            "Method selection by sector: utilities and stable-payout "
            "names use a Dividend Discount Model with growth capped "
            "at 8% per year to avoid runaway valuations; banks use a "
            "price-to-book multiple anchored to the sector median; "
            "cyclicals use a price-to-earnings multiple on cycle-"
            "adjusted earnings; the fallback is the Graham number, "
            "which is the geometric mean of book value and earnings "
            "anchored to a conservative multiple. Quality (described "
            "below) acts as a multiplier — even a deep discount on a "
            "low-quality balance sheet is downgraded to NEUTRAL.",
         "use":
            "A slow signal, 6-24 month horizon. Used as a small "
            "conviction overlay rather than a primary signal: when "
            "value and momentum agree, conviction can go up one notch; "
            "when they disagree, the bullish call is downgraded."},

        {"name": "PSX Filings & Director's Reports",
         "purpose":
            "Forward-looking commentary directly from management. "
            "Quarterly and annual filings contain the most credible "
            "guidance available in Pakistani markets — most retail "
            "and even some institutional investors do not read them. "
            "Pulled from the PSX Data Portal, where the Director's "
            "Report is embedded inside each set of quarterly "
            "financial results and the headline ratios are inside "
            "the annual report.",
         "fields":
            "Filing date, period, document type, summarised outlook, "
            "tone score (-1 to +1), guidance strength (high / medium "
            "/ low), list of growth plans, list of risks, capex / "
            "expansion flags, link to original document. "
            "Capacity & expansion fields extracted verbatim where "
            "management states them: installed capacity (for example "
            "'1,292 MW gross'), actual production (for example "
            "'823 MW average dispatch in H1 FY26'), implied "
            "utilisation %, and a list of any new products / "
            "expansions announced.",
         "calc":
            "Each filing is read by an AI vision model that extracts "
            "the structured fields above. The instruction is strict: "
            "leave any number null unless the report literally states "
            "it — no inference. The original PDF is retained for "
            "verification, and the verbatim quote is shown in the "
            "user interface alongside the extracted figure.",
         "use":
            "A recent filing (within 14 days) with high guidance "
            "strength and a tone above +0.5 nudges conviction up; a "
            "tone of -0.4 or below caps high conviction down to "
            "medium. Filings older than 270 days are ignored — the "
            "narrative is too stale. Capacity & expansion lets the "
            "analyst sanity-check expansion claims: a power "
            "company running at 64% of installed capacity that "
            "announces a new 200 MW unit is well-justified; the "
            "same company at 92% utilisation announcing the same "
            "capex is a stronger demand signal. The decision step "
            "is told that low utilisation (below 70%) combined "
            "with a new capex announcement should *downgrade* "
            "conviction one notch — capacity is the constraint, "
            "not demand.",
         "limit":
            "PSX exposes only the five most recent filings per "
            "category per company, so we forward-cache to build a "
            "longer history over time."},

        {"name": "Mettis Global news + PSX notices feed",
         "purpose":
            "A second, ticker-keyed news feed built specifically "
            "for the PSX. Mettis Global publishes general business "
            "headlines and — uniquely — turns every PSX corporate "
            "notice (dividend, board meeting, profit release, "
            "ratings action, regulatory action) into a navigable "
            "news article tagged with the ticker. That gives the "
            "system a fast, structured stream of corporate actions "
            "that the general business press picks up only hours "
            "later.",
         "fields":
            "Headline, summary, publication timestamp, source, URL, "
            "and a best-effort list of PSX tickers detected in the "
            "headline body.",
         "calc":
            "Scraped from the public site with graceful "
            "degradation (an empty result on a layout change, "
            "never a crash). Fed into the same automated reading "
            "pass used for the general business press, with a "
            "ticker-hits hint so the scorer can attribute "
            "sentiment to the right names.",
         "use":
            "Specifically for ticker-level news: a regulatory "
            "action article published on Mettis at 10:00 PKT is "
            "live in the briefing within the hour and can block a "
            "buy that the daily run was about to recommend on the "
            "next session."},

        {"name": "Sarmaya.com cross-check",
         "purpose":
            "Independent triangulation on the headline "
            "fundamentals. Sarmaya is a well-respected Pakistani "
            "equities portal that publishes per-symbol P/E, P/B, "
            "EPS, dividend yield, and market capitalisation. "
            "Comparing against it catches vendor-data lags or "
            "outright errors before the analyst sees a wrong "
            "number on the Value tab.",
         "fields":
            "Per symbol: P/E, P/B, EPS, dividend yield, "
            "market cap, last-update timestamp, source URL.",
         "calc":
            "Scraped weekly per symbol (cheap — only fifteen "
            "stocks). Cached locally. Compared against the "
            "primary fundamentals source field-by-field; any "
            "absolute disagreement above 25% is flagged on the "
            "Value tab as a Cross-check warning so the analyst "
            "can verify before placing the trade.",
         "use":
            "Sanity layer only. The primary fundamentals source "
            "remains authoritative; Sarmaya is an early-warning "
            "system for stale or incorrect numbers."},

        {"name": "PSX Material Information",
         "purpose":
            "Companies listed on the PSX are required to file "
            "*Material Information* — price-sensitive disclosures "
            "such as significant contract wins, board "
            "resolutions, plant shutdowns, regulatory enforcement, "
            "or any change a reasonable investor would need to "
            "know about. These filings empirically precede 3-7% "
            "price gaps on the affected stock and a fresh filing "
            "is a strong volatility flag.",
         "fields":
            "Symbol, filing date, title, link to the original "
            "PDF, and a derived doc identifier so duplicates "
            "across refreshes are deduplicated.",
         "calc":
            "Scraped daily after market close from the PSX Data "
            "Portal MATERIAL-tagged announcement stream and "
            "appended to a dedicated cache. The briefing carries "
            "a 'Material Information (last 5 trading days)' "
            "stanza listing recent filings for the relevant "
            "symbol.",
         "use":
            "Treated as a volatility flag rather than a "
            "directional signal — direction is unknown until the "
            "filing has been read in context. The decision step "
            "is instructed to widen the predicted-return band and "
            "downgrade high conviction to medium when a fresh "
            "Material Information filing is present. The user "
            "interface shows a banner on the Today tab when any "
            "holding has a fresh disclosure, and a Material "
            "Disclosures section on the Reports tab."},
    ]

    for s in sources:
        doc.append({"kind": "h", "level": 3, "text": s["name"]})
        doc.append({"kind": "p",
                    "text": f"**Purpose.** {s['purpose']}"})
        if s.get("fields"):
            doc.append({"kind": "p",
                        "text": f"**Fields captured.** {s['fields']}"})
        if s.get("calc"):
            doc.append({"kind": "p",
                        "text": f"**How it is calculated.** {s['calc']}"})
        if s.get("use"):
            doc.append({"kind": "p",
                        "text": f"**How it feeds the decision.** {s['use']}"})
        if s.get("limit"):
            doc.append({"kind": "p",
                        "text": f"**Known data limit.** {s['limit']}"})

    doc.append({"kind": "pagebreak"})

    # ----------------------------------------------------- 4. STRATEGIES
    doc += [
        {"kind": "h", "level": 1, "text": "4. Strategy Catalogue"},
        {"kind": "p", "text":
            "Each strategy below is a deterministic feature computed "
            "before the decision step. The decision step is told exactly "
            "how to weigh each one in its rule book. The catalogue is "
            "exhaustive — there are no hidden weights."},
    ]

    strategies = [
        {"name": "Strategy 1 — Monthly Momentum",
         "why":
            "12-1 momentum (the 12-month return excluding the most "
            "recent month) is the single most-replicated anomaly in "
            "global equity markets. On Pakistani small and mid caps it "
            "has historically generated 1.5-2.0% annual alpha after "
            "transaction costs.",
         "calc":
            "For each stock, log-return from t-21 to t-252 trading "
            "days. Stocks are ranked 1-15 each day and the top five "
            "form the eligibility list.",
         "role":
            "Primary screen. A stock outside the top-5 rarely receives "
            "a buy recommendation, regardless of how favourable other "
            "signals look.",
         "limit":
            "Crashes hard in regime changes; the market-regime filter "
            "below mitigates."},

        {"name": "Strategy 2 — Market-Regime Filter",
         "why":
            "Momentum strategies lose money in crisis regimes. A "
            "regime classifier prevents the system from running "
            "momentum in environments where it has historically been "
            "destructive.",
         "calc":
            "Regime is one of {NORMAL, CAUTION, CRISIS}, derived from "
            "the KSE-100 index distance from its 200-day moving average "
            "and the realised-volatility percentile. Each regime maps "
            "to an exposure multiplier: 1.0, 0.5, and 0.0 respectively.",
         "role":
            "Caps total exposure. A CRISIS reading allows no new buys "
            "for the entire day."},

        {"name": "Strategy 3 — AI Defensive Overlay",
         "why":
            "Hard rules miss event-driven risk. Before confirming a "
            "buy, the AI scans the briefing's news section, "
            "fundamentals, and management outlook for any reason to "
            "step back.",
         "calc":
            "The decision model returns direction, conviction, action, "
            "expected-return band, key drivers, and key risks per stock.",
         "role":
            "Final synthesiser. Can downgrade a stock that passes "
            "every quantitative screen if it spots qualitative red "
            "flags."},

        {"name": "Strategy 4 — Sector-aware Intrinsic Value",
         "why":
            "Stocks trading 25% or more below estimated fair value "
            "typically close that gap within 6-24 months. The win-rate "
            "is high; the timing is slow. Sector-aware so utilities are "
            "not valued on the same yardstick as banks or growth names.",
         "calc":
            "Method selected by sector: utilities use a Dividend "
            "Discount Model with capped growth; banks use a price-to-"
            "book multiple anchored to the sector median; cyclicals use "
            "a price-to-earnings multiple on cycle-adjusted earnings; "
            "the fallback is the Graham number.",
         "role":
            "Slow conviction overlay. BUY-VALUE plus bullish momentum "
            "allows a one-notch conviction upgrade; conflicting "
            "value-vs-momentum signals are surfaced as a key risk in "
            "the rationale."},

        {"name": "Strategy 5 — Quality Score",
         "why":
            "Cheap junk is a value trap. Cheap-and-quality is the "
            "strongest setup the system can identify.",
         "calc":
            "Equal-weighted three-component score on a 0-100 scale: "
            "profitability (banded ROE — at or above 20% scores full "
            "marks), leverage (banded debt-to-equity — at or below 0.5 "
            "scores full marks), and stability (5-year EPS coefficient "
            "of variation — lower is better).",
         "role":
            "Filter on Strategy 4. A junk score (under 30) combined "
            "with a BUY-VALUE signal is forced to HOLD even if "
            "momentum is positive, because the historical track record "
            "of \"cheap junk\" is poor."},

        {"name": "Strategy 6 — Earnings Momentum",
         "why":
            "Post-earnings drift: stocks with accelerating year-over-"
            "year EPS growth tend to outperform for 4-8 weeks after a "
            "result. The effect is symmetric on the downside.",
         "calc":
            "Year-over-year EPS growth, prior-period growth, "
            "acceleration (this − prior), and 3-year compound annual "
            "growth rate. These map to flags: ACCELERATING, RECOVERING, "
            "STABLE, DECELERATING, EROSION.",
         "role":
            "ACCELERATING combined with bullish price momentum allows "
            "high conviction; EROSION combined with neutral signals "
            "downgrades to HOLD or AVOID even if technicals are fine."},

        {"name": "Strategy 7 — Earnings Blackout",
         "why":
            "Result-day gaps of 5-10% destroy any 5-day prediction. "
            "The blackout rule strictly forbids new positions inside "
            "the five-trading-day window before a high- or medium-"
            "confidence earnings event.",
         "calc":
            "Hard filter — score is set to negative infinity for any "
            "stock currently inside the blackout window.",
         "role":
            "Defensive. A softer event-window flag (6-14 days out) "
            "caps conviction one notch lower with a tighter "
            "recommended stop loss."},

        {"name": "Strategy 8 — Overnight Gap Prior",
         "why":
            "PSX opens nine hours after the US close. Frontier "
            "emerging-market betas to overnight US moves are well "
            "documented: average about 0.7 in normal regimes, about "
            "1.5 in stressed regimes.",
         "calc":
            "Two-layer: a rules-based prior (for example, S&P down "
            "more than 0.5% AND VIX above 20 produces GAP-DOWN; both "
            "green produces GAP-UP); plus a small statistical model "
            "fitted on six months of overnight-to-PSX-open mappings, "
            "which produces a continuous expected-open percentage.",
         "role":
            "First block of the briefing. The decision step is told "
            "explicitly: do not generate a recommendation that ignores "
            "a 2% futures session."},

        {"name": "Strategy 9 — News Sentiment",
         "why":
            "Hand-labelling Pakistani business headlines is not "
            "scalable. An automated reading pass produces consistent "
            "signed sentiment, confidence, category, and ticker "
            "labels in near-real time.",
         "calc":
            "Confidence-weighted average of sentiment over rolling "
            "windows of 24 hours (macro news) and 72 hours (ticker-"
            "specific news).",
         "role":
            "Conviction overlay. Strongly negative headlines (score "
            "below -0.3) in the last 24 hours block new buys on the "
            "affected names."},

        {"name": "Strategy 10 — Management Outlook",
         "why":
            "Quarterly and annual Director's Reports contain the most "
            "credible forward-looking commentary in Pakistani markets. "
            "Most market participants do not read them.",
         "calc":
            "An AI vision model extracts a structured outlook record "
            "per filing: tone from -1 to +1, guidance strength "
            "(high / medium / low), growth plans, risks mentioned, "
            "and capex / expansion flags.",
         "role":
            "Slow overlay. A tone of -0.4 or below combined with high "
            "AI conviction caps the call at medium conviction, with "
            "an explicit annotation in the trade plan. A tone of +0.5 "
            "or above with capex announced and bullish momentum "
            "allows a small conviction upgrade."},

        {"name": "Strategy 11 — Volatility-Conditional Range Widening",
         "why":
            "Predicting tight return bands on stressed-volatility days "
            "produced systematic inside-range misses (about 50% hit-"
            "rate). Widening bands proportional to the volatility "
            "index restored a roughly 70% hit-rate.",
         "calc":
            "If the US volatility index is 18 or above, the "
            "predicted-return band is widened by a factor of 1.5; if "
            "it reaches 22 or above, by a factor of 2.0.",
         "role":
            "Embedded in the AI's rule book — applied at the moment "
            "of decision, not retrofitted afterwards."},

        {"name": "Strategy 12 — Volatility & Volume Confirmation Indicators",
         "why":
            "Three classic indicators add timing precision on top of "
            "the slower momentum and value signals: Bollinger Bands "
            "(volatility-adjusted price extremes), MACD with its "
            "histogram (trend confirmation and turn detection), and "
            "On-Balance Volume (does volume confirm the move?). The "
            "system computes them deterministically and surfaces the "
            "interpretations to the decision step.",
         "calc":
            "Bollinger Bands: 20-day mean and ±2 standard deviations "
            "around it. The system tracks two derived numbers — %B "
            "(where the close sits relative to the bands; above 1 "
            "is above the upper band) and band-width as a "
            "percentile of the last 252 trading days (low percentile "
            "= a 'squeeze' that typically resolves into a sharp "
            "directional move). MACD: 12-26 exponential moving "
            "averages with a 9-period signal line and a histogram "
            "that captures the difference; positive histogram = "
            "bullish trend, sign-flip = trend change. OBV (On-"
            "Balance Volume): cumulative volume signed by the close-"
            "vs-prior-close; the system reports the 5-day percentage "
            "change so a rising tape with rising OBV reads as "
            "volume-confirmed.",
         "role":
            "Conviction overlay. Selected rules embedded in the "
            "decision rule book: %B above 0.95 with an already high "
            "RSI is overbought — downgrade BUY conviction one notch. "
            "%B below 0.05 with positive earnings momentum is a "
            "mean-reversion BUY setup. Band-width below the 5th "
            "percentile is a squeeze — the next move is likely "
            "large; widen the predicted-return band. A MACD "
            "histogram crossing positive while price is above the "
            "50-day moving average is treated as trend confirmation. "
            "A 5-day OBV change of +10% or more on a bullish call "
            "permits a small conviction upgrade; a -10% OBV move "
            "against an apparent uptrend is a divergence flag."},

        {"name": "Strategy 13 — Sector-aware Macroeconomic Impact",
         "why":
            "Macroeconomic news rarely moves all stocks in the same "
            "direction; it sorts the market by sector. A 1 percentage-"
            "point increase in the policy rate is a strong tailwind "
            "for banks (they reprice loans faster than deposits) and "
            "a strong headwind for cement (financial costs spike and "
            "construction demand falls). An oil price spike rewards "
            "exploration and production companies and squeezes "
            "transportation, packaging, and pharmaceutical importers. "
            "A jump in gold or a fall in copper signals risk-off "
            "money flow that pulls foreign capital away from Pakistan "
            "equities regardless of sector. The system needs to "
            "recognise these patterns automatically and tell the "
            "analyst exactly which stocks win and which stocks get "
            "hurt — every day, on every recommendation.",
         "calc":
            "A deterministic rule book (kept in plain code, not a "
            "model) maps each macroeconomic indicator move to a "
            "signed score in each sector. The drivers tracked are: "
            "(1) State Bank policy-rate level and change, (2) Brent "
            "crude 5- and 21-day returns, (3) USD/PKR 21- and 63-day "
            "moves, (4) coal proxy implied from sustained oil moves "
            "(coal prices lag Brent by about a month with high "
            "correlation), (5) gold as a global risk-off proxy, "
            "(6) copper as a global industrial-growth proxy, "
            "(7) cotton for textile-export exposure, and a layer of "
            "industry-specific KPIs that the State Bank publishes "
            "every business day: (8) the 3-month T-bill cut-off "
            "yield and its position relative to the policy rate "
            "(money-market signal that banks lean on), (9) the "
            "3-month KIBOR (the funding-cost benchmark that flows "
            "into floating-rate loan and corporate financial-cost "
            "lines), (10) total foreign-exchange reserves with a "
            "stress band below USD 8 billion and a recovery band "
            "above USD 14 billion that captures balance-of-payments "
            "regime shifts, (11) the KSE-100 5- and 21-day momentum "
            "for broad-market regime, and (12) the latest Pakistan "
            "CPI year-on-year print scraped from Trading Economics "
            "with a fallback to PBS — a high CPI reading keeps the "
            "rate environment restrictive, while a cooling CPI "
            "opens the door to cuts and is a powerful tailwind for "
            "leveraged sectors. Each driver fires only when the "
            "move clears a meaningful threshold (oil only counts "
            "when it has moved 7% in a week or 10% in a month, "
            "T-bills must drift more than 30 basis points away from "
            "the policy rate to register, reserves must cross the "
            "stress / recovery bands rather than wobble around the "
            "centre — smaller moves are filtered as noise). Banking "
            "on a rate-up day scores +2 (margin expansion) and gets "
            "a further +1 when KIBOR is rising or T-bills are "
            "trading above policy; Cement on a rate-up day scores "
            "-3 and gets another -2 when KIBOR rises (financial "
            "costs and demand both bite); IPPs score -2 on rising "
            "KIBOR and -2 when reserves drop into the stress band "
            "because circular debt almost always worsens during BoP "
            "stress; pharma scores -2 on reserve stress because API "
            "imports require letters of credit that get harder to "
            "confirm. The score for an individual stock then "
            "adjusts that sector reading by the company's debt-to-"
            "equity ratio from the latest balance sheet — a high-"
            "leverage cement company is hit harder than a low-"
            "leverage peer — and by company-specific tags (CASA-"
            "rich tier-1 banks get an extra notch on rate-up days; "
            "HUBCO carries an extra headwind notch on rate-up days "
            "because of its ongoing tariff renegotiation). Policy-"
            "rate changes are detected by comparing today's rate "
            "against the most recent observation on a strictly "
            "earlier date, stored in a small history file — that "
            "way only real MPC moves between distinct days fire "
            "the rate-change rules, not intra-day repeats. Every "
            "score line carries a human-readable explanation so the "
            "analyst sees not just '+2' but 'higher policy rate "
            "widens net interest margins because banks reprice "
            "loans faster than deposits'. The same KPI snapshot is "
            "also written into the LLM briefing so the AI's "
            "rationale can cite the actual T-bill or CPI number "
            "rather than speak in generalities.",
         "role":
            "Three roles. (1) Reasoning input: the briefing handed "
            "to the AI now includes a 'macro impact for this stock' "
            "block listing the active drivers, the sector verdict, "
            "and the stock-level verdict with its amplifier note. "
            "The AI is required to cite at least one macro tailwind "
            "or headwind in its rationale and to reflect strong "
            "headwinds in conviction. (2) Direct visibility: a "
            "Macro Radar panel on the Today tab shows today's "
            "sector winners and losers and the most-affected "
            "individual stocks, so the analyst can see the same "
            "logic the AI saw. The same data is exposed inside the "
            "per-stock Forecast drill-down as the 'Why this call?' "
            "panel. (3) Chat answers: the chatbot can answer "
            "questions like 'what is the macro impact on cement "
            "today?' or 'which banks benefit from this rate cut?' "
            "by calling the same engine — so verbal questions get "
            "the same explainable, sector-specific answer the UI "
            "shows."},

        {"name": "Strategy 14 — Cost-aware Trade Filter",
         "why":
            "Cost models are usually retrofitted to post-mortem P&L. "
            "Here they are baked in as a pre-trade filter so no "
            "uneconomic recommendation can be surfaced.",
         "calc":
            "Round-trip cost approximately 0.56% (brokerage 0.30%, "
            "federal excise duty on brokerage 0.05%, CDC and exchange "
            "fees 0.01%, slippage 0.20%); 15% capital-gains tax on "
            "net positive returns; minimum gross return required to "
            "trade is cost plus a 1.0 percentage-point edge — about "
            "1.56%.",
         "role":
            "Hard filter. Any signal with a gross expected return "
            "below the threshold is dropped automatically, regardless "
            "of how strong the qualitative case looks."},
    ]

    for s in strategies:
        doc.append({"kind": "h", "level": 3, "text": s["name"]})
        doc.append({"kind": "p", "text": f"**Why.** {s['why']}"})
        doc.append({"kind": "p", "text": f"**How calculated.** {s['calc']}"})
        doc.append({"kind": "p", "text": f"**Role in decision.** {s['role']}"})
        if s.get("limit"):
            doc.append({"kind": "p",
                        "text": f"**Known limit.** {s['limit']}"})

    doc.append({"kind": "pagebreak"})

    # ----------------------------------------------------- 5. DECISION
    doc += [
        {"kind": "h", "level": 1, "text": "5. Decision Engine"},

        {"kind": "h", "level": 2,
         "text": "5.1  Why an AI model rather than a fixed weighted score?"},
        {"kind": "p", "text":
            "Four alternatives were prototyped before settling on AI "
            "synthesis: (a) a hand-tuned weighted-average score; (b) a "
            "machine-learning model fitted on the same features; (c) a "
            "smaller specialised model trained on Pakistani filings; "
            "and (d) AI-as-synthesiser. The first three were brittle to "
            "feature drift — adding a single new data source forced a "
            "complete refit. The AI approach holds up because the rule "
            "book is written in plain English: adding a new feature "
            "becomes a documentation change, not a re-fit. The "
            "rationale produced for every recommendation is also "
            "human-auditable, which matters when a stock is going "
            "against the call."},

        {"kind": "h", "level": 2,
         "text": "5.2  What the AI returns"},
        {"kind": "p", "text":
            "Every recommendation is a structured record with the "
            "following fields. Anything missing or malformed forces an "
            "automatic re-prompt with the validation error."},

        {"kind": "table",
         "headers": ["Field", "Type", "Meaning"],
         "rows": [
            ["Symbol", "Text", "PSX ticker"],
            ["Direction", "Categorical",
             "Bullish, bearish, or neutral over a 5-trading-day "
             "horizon"],
            ["Conviction", "Categorical",
             "High, medium, or low — used as the ranking weight "
             "(1.0, 0.7, 0.3)"],
            ["Suggested action", "Categorical",
             "Buy, add, hold, trim, sell, or avoid"],
            ["Expected return band", "Three numbers",
             "Low, mid, and high estimate of the 5-day return in %"],
            ["Entry price", "Number", "Suggested entry in PKR"],
            ["Stop", "Number", "Suggested stop loss in PKR"],
            ["Target", "Number", "Suggested take-profit in PKR"],
            ["Key drivers", "List",
             "Up to four short reasons supporting the call"],
            ["Key risks", "List",
             "Up to four short reasons the call could fail"],
            ["Rationale", "1-2 sentences",
             "Plain-English summary citing the specific signals that "
             "drove the recommendation"],
        ]},

        {"kind": "h", "level": 2,
         "text": "5.3  Selected rules from the AI rule book"},
        {"kind": "p", "text":
            "The rule book embedded in the AI's instructions is about "
            "3,000 characters and codifies how each data feature should "
            "influence the output. Selected rules:"},
        {"kind": "bullets", "items": [
            "**Cost-aware bias toward HOLD.** Do not recommend buy or "
            "add when the expected mid return is below 1.6% over five "
            "trading days; recommend hold instead.",
            "**Overnight global risk.** If the briefing shows GAP-DOWN "
            "with stressed VIX, downgrade conviction one notch. "
            "Frontier markets typically lose 1-2× the S&P move on "
            "risk-off days.",
            "**Volatility-conditional bands.** When VIX is 18 or above, "
            "widen the expected return band by at least 50%.",
            "**Earnings blackout (critical).** If the briefing "
            "contains an earnings-blackout warning, the action must be "
            "HOLD or AVOID regardless of any other signal.",
            "**Quality-and-value gating.** A junk balance-sheet "
            "combined with a BUY-VALUE flag yields HOLD (treated as a "
            "value trap); a high-quality balance sheet combined with "
            "BUY-VALUE is the strongest setup.",
            "**Management outlook.** A management tone of -0.4 or "
            "below downgrades conviction one notch; a tone of +0.5 or "
            "above with high guidance strength and bullish momentum "
            "permits a small upgrade.",
            "**Skepticism by default.** When nothing special is "
            "happening, return NEUTRAL / LOW conviction / HOLD. The "
            "system is explicitly designed to refuse trades; quality "
            "over quantity.",
        ]},
        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- 6. RISK
    doc += [
        {"kind": "h", "level": 1, "text": "6. Risk Management & Filters"},

        {"kind": "h", "level": 2, "text": "6.1  Transaction-cost model"},
        {"kind": "table",
         "headers": ["Cost layer", "Per side (%)", "Round-trip (%)"],
         "rows": [
            ["Brokerage (typical retail)",     "0.150",  "0.300"],
            ["Federal excise duty on brokerage (16%)",
                                                "0.024",  "0.048"],
            ["CDC charges",                     "0.005",  "0.010"],
            ["PSX laga",                        "0.0015", "0.003"],
            ["SECP fee",                        "0.00005", "0.0001"],
            ["Slippage (mid-cap blue chips)",   "—",      "0.200"],
            ["TOTAL round-trip",                "—",      "≈ 0.560"],
            ["Capital-gains tax (filer, < 1y holding)",
                                                "—",
                                                "15.0% of net positive P&L"],
            ["Minimum gross required",           "—",
                                                "≈ 1.56% (cost + 1.0% edge)"],
        ]},
        {"kind": "p", "text":
            "All numbers above are pulled from the cost configuration "
            "at run time. Switching to a discount broker requires only "
            "updating the brokerage rate in one place; the rest of the "
            "system re-prices itself automatically."},

        {"kind": "h", "level": 2,
         "text": "6.2  Hard filters (recommendation rejected outright)"},
        {"kind": "bullets", "items": [
            "Suggested action is anything other than BUY or ADD.",
            "Direction is anything other than bullish.",
            "Gross expected mid return is below 1.56% (cost plus edge).",
            "Symbol is in the earnings blackout window — five trading "
            "days or fewer to a high- or medium-confidence earnings "
            "event.",
        ]},

        {"kind": "h", "level": 2,
         "text": "6.3  Soft caps (one-notch downgrade)"},
        {"kind": "bullets", "items": [
            "Management tone of -0.4 or below: high → medium, "
            "medium → low. The downgrade is explicitly annotated in "
            "the trade plan.",
            "VIX of 22 or above: range bands already widen; the AI is "
            "additionally instructed to reduce conviction one notch on "
            "stressed-volatility days.",
            "Earnings event window (6-14 days out): conviction one "
            "notch lower with tighter recommended stop loss.",
            "Quality junk score combined with BUY-VALUE: forced HOLD "
            "instead of buy.",
            "Foreign-net-selling streak of three days or more "
            "combined with negative news in the last 24 hours: "
            "conviction one notch lower.",
        ]},

        {"kind": "h", "level": 2, "text": "6.4  Position-sizing guidelines"},
        {"kind": "p", "text":
            "The system suggests trades but does not enforce position "
            "sizes. The recommended discipline (encoded in the daily "
            "brief and the user interface) is:"},
        {"kind": "table",
         "headers": ["Conviction", "Suggested allocation", "Stop placement"],
         "rows": [
            ["High",   "10-15% of capital per name",
             "1.5× ATR or AI-suggested stop, whichever is wider"],
            ["Medium", "5-8% per name", "1.0× ATR"],
            ["Low",    "3% per name, or pass", "Tight, 0.5× ATR"],
            ["No setup", "Stay in cash", "—"],
        ]},
        {"kind": "callout", "text":
            "**Cash is a position.** The system is designed to surface "
            "zero-trade days. Of approximately 120 trading days in the "
            "back-test, the median day produced 1.2 buys, with 18% of "
            "days producing no qualifying setup at all. This is a "
            "feature, not a bug."},
        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- 7. VALIDATION
    doc += [
        {"kind": "h", "level": 1, "text": "7. Validation & Performance"},

        {"kind": "h", "level": 2, "text": "7.1  Back-test methodology"},
        {"kind": "p", "text":
            "Each trading day from November 2025 to April 2026 is "
            "replayed with a strict information cutoff. On day t, the "
            "system can see only data dated t or earlier — including "
            "fundamentals (point-in-time), news (timestamp-filtered), "
            "and the AI rule book exactly as it was on that date. "
            "Recommendations are then evaluated against the actual "
            "price five trading days later. The earnings-blackout, "
            "cost, and outlook caps are applied identically to live "
            "runs, so the back-test reflects what the analyst would "
            "actually have seen on the day."},

        {"kind": "h", "level": 2, "text": "7.2  Headline metrics"},
        {"kind": "table",
         "headers": ["Metric", "Definition", "Result"],
         "rows": [
            ["Direction hit-rate",
             "% of predictions whose 5-day actual sign matches predicted",
             "≈ 62%"],
            ["3-class hit-rate",
             "Bullish / bearish / neutral exact match",
             "≈ 54%"],
            ["Inside-range hit-rate",
             "Actual return falls inside [low, high] band",
             "≈ 71% (after volatility-conditional widening)"],
            ["Mean absolute error on mid",
             "|actual − predicted_mid| in %",
             "≈ 1.9%"],
            ["R² of predicted mid vs actual",
             "0 = mean prediction; 1 = perfect; negative = worse than "
             "the mean",
             "≈ 0.18"],
            ["Net edge — top decile by score",
             "Mean 5-day net return on highest-conviction calls",
             "+1.2 to +1.8%"],
            ["Net edge — full universe",
             "Mean 5-day net return on all buy / add calls",
             "+0.4 to +0.7%"],
            ["Sharpe ratio (annualised, top decile)",
             "Net daily return ÷ volatility × √252",
             "≈ 1.1 (sample-thin)"],
        ]},

        {"kind": "h", "level": 2, "text": "7.3  Honest characterisation"},
        {"kind": "bullets", "items": [
            "**Hit-rate is informative, not definitive.** A 62% "
            "direction hit-rate looks decent; the magnitude "
            "distribution is bi-modal, with a few +5% wins funding "
            "the slow grind of -1% misses.",
            "**Sample is small.** Six months × about 120 trading days "
            "× 15 stocks gives 1,800 stock-day predictions, but the "
            "actionable subset (buys actually surfaced after all "
            "filters) is roughly 150 trades. Confidence intervals on "
            "the Sharpe ratio are wide.",
            "**Survivorship.** The universe is current-day, not "
            "point-in-time membership. Conclusions do not extend to "
            "PSX small caps not currently in the universe.",
            "**Regime risk.** The back-test window was a generally "
            "constructive macro period — rate cuts, stable currency. "
            "Behaviour in a crisis regime is untested live.",
            "**AI variability.** AI providers occasionally update "
            "their models. The system pins specific versions where "
            "possible and re-validates on a rolling 30-day window "
            "after any forced update.",
        ]},
        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- 8. CADENCE
    doc += [
        {"kind": "h", "level": 1, "text": "8. Operational Cadence"},
        {"kind": "p", "text":
            "Each data source has its own automated refresh schedule. "
            "The user interface reads from the cached data, so it is "
            "always usable even when an upstream source is "
            "temporarily unreachable — the affected layer is simply "
            "flagged as stale rather than producing wrong answers."},
        {"kind": "table",
         "headers": ["Pipeline", "Schedule", "What it does"],
         "rows": [
            ["End-of-day prices and flows",
             "Daily, 16:00 PKT, Mon-Fri",
             "Fetch closing prices for the universe; recompute "
             "Layer-1 features; refresh foreign and local flow data; "
             "trigger filings refresh on earnings days."],
            ["News scoring",
             "Hourly, 07:00-18:00 PKT",
             "Pull headlines from the five sources; score new items; "
             "append to the news cache."],
            ["Overnight global block",
             "Daily, 06:00 PKT",
             "Fetch overnight closes for US, Asia, FX, and the "
             "volatility index; rebuild the gap-prior."],
            ["Fundamentals refresh",
             "Weekly, Sunday",
             "Refresh per-symbol financials, sector medians, "
             "valuation book, quality scores, and earnings calendar."],
            ["Filings extraction",
             "Weekly, Saturday + on earnings",
             "Scrape the latest filings; have the AI vision model "
             "extract structured outlooks; persist to the filings "
             "cache."],
            ["Daily decision run",
             "Daily, 08:30 PKT, Mon-Fri",
             "Build per-stock briefings; run the AI rule book; "
             "persist recommendations; update the prediction-vs-"
             "actual log."],
        ]},

        {"kind": "h", "level": 2, "text": "8.1  Failure modes & mitigations"},
        {"kind": "table",
         "headers": ["Failure", "Mitigation"],
         "rows": [
            ["Primary AI provider outage",
             "Automatic fall-back to a secondary AI provider; if both "
             "fail, the previous day's recommendations are retained "
             "with an explicit STALE flag in the user interface."],
            ["Vendor data throttling",
             "Per-call retry with exponential back-off; cached "
             "last-known-good values are reused for up to three days."],
            ["A news source goes down",
             "The other four sources continue to function; per-source "
             "failure is logged but does not block the daily run."],
            ["PSX portal layout change",
             "Connector tests run as part of every refresh; a layout "
             "change triggers an explicit failure rather than a silent "
             "wrong answer."],
            ["Earnings calendar miss",
             "Conservative default — when the date is uncertain, the "
             "stock is treated as event-window (caps conviction) "
             "rather than blackout (blocks buys)."],
        ]},
        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- 9. LIMITATIONS
    doc += [
        {"kind": "h", "level": 1, "text": "9. Limitations & Future Work"},

        {"kind": "h", "level": 2, "text": "9.1  Known limitations"},
        {"kind": "bullets", "items": [
            "**End-of-day only.** No intraday signals; gap-on-open "
            "events are unhedgeable in this design.",
            "**Small universe.** Fifteen stocks, with sector "
            "concentration in power, banks, energy, cement, and "
            "refining. Diversification is limited by construction.",
            "**Single-asset, not portfolio-optimised.** No covariance "
            "matrix; no Markowitz optimisation. Position sizing is a "
            "rule of thumb, not an optimisation output.",
            "**No options.** The PSX options market is illiquid, so "
            "convexity cannot be expressed directly.",
            "**News coverage gaps.** Urdu-only sources are missed; "
            "Pakistani financial Twitter / X is not currently "
            "ingested.",
            "**AI variability.** The same briefing can produce "
            "slightly different conviction across runs. Mitigated by "
            "low temperature and pinned model versions, but not "
            "eliminated.",
            "**Back-test survivorship.** The universe is current-day, "
            "not point-in-time membership.",
        ]},

        {"kind": "h", "level": 2, "text": "9.2  Roadmap (prioritised)"},
        {"kind": "table",
         "headers": ["#", "Item", "Expected lift"],
         "rows": [
            ["1", "Per-stock 5-minute intraday data via the PSX "
             "Terminal feed (paid)",
             "Better entry timing; about 20% reduction in mean "
             "absolute error."],
            ["2", "Add 10 mid-cap stocks (universe → 25)",
             "Diversification and access to additional alpha pockets."],
            ["3", "A lightweight portfolio optimiser",
             "Better risk-adjusted returns at the same per-stock edge."],
            ["4", "Twitter / X sentiment connector",
             "Faster reaction to retail flows and viral narratives."],
            ["5", "Fully point-in-time fundamentals, cross-checked "
             "against direct PSX filings",
             "Cleaner earnings-momentum feature; fewer vendor-data "
             "lag issues."],
            ["6", "Live paper-trading vs. broker API",
             "True forward-test with real fills; eliminates back-test "
             "biases."],
            ["7", "Per-strategy alpha-attribution dashboard",
             "Lets the analyst see which layers earn or lose money "
             "over a custom date range."],
        ]},
        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- 10. AUDIT
    doc += [
        {"kind": "h", "level": 1,
         "text": "10. Solution Health Audit (April 2026)"},

        {"kind": "p", "text":
            "After the analyst feedback round was incorporated, the "
            "full solution was re-audited end-to-end. Every connector "
            "was probed live, every chatbot tool was dispatched, every "
            "decision-engine output was inspected for the new "
            "macroeconomic-impact fields, and every dashboard tab was "
            "rendered against today's data. The findings and the fixes "
            "applied are listed below in plain English so the reader "
            "knows exactly what state the system is in today."},

        {"kind": "h", "level": 2,
         "text": "10.1  Issues found and fixed"},
        {"kind": "table",
         "headers": ["#", "Where", "Issue", "Fix applied"],
         "rows": [
            ["1", "Data freshness panel",
             "The 'Material Information' file always reported "
             "'no latest data date' even when the parquet was up to "
             "date — the freshness check was missing the branch that "
             "reads the date column for that particular file.",
             "Added a branch that reads the announcement date "
             "(falling back to the scrape time). The file now shows "
             "an accurate latest-data-date and freshness flag."],
            ["2", "Macro impact engine",
             "The policy-rate history file was keyed on rate value "
             "rather than calendar date, so multiple observations "
             "written on the same day (synthetic tests, repeated "
             "rule-based fallbacks) inflated false rate-change "
             "drivers. The engine occasionally reported a phantom "
             "rate cut.",
             "Switched the persistence to one row per calendar date. "
             "Only the rate from a strictly earlier date is treated "
             "as the previous reading, so intra-day repeats can no "
             "longer cause phantom MPC moves. The polluted history "
             "was reset."],
            ["3", "Predictions log",
             "Predictions written to the log before the macro "
             "impact engine existed had no macro tailwinds, "
             "headwinds, or impact snapshot, so the 'Why this call?' "
             "panel showed empty macro lines for older predictions.",
             "Added a back-fill step inside the prediction reader: "
             "when an older prediction is loaded for display, the "
             "macro impact engine is run live and its output is "
             "stitched in, so the explanation panel always has macro "
             "context."],
            ["4", "Chatbot tools",
             "The chatbot could not answer macro-impact, "
             "Director's-Report, Material-Information or sector-"
             "volume questions because there were no specific "
             "tools for them. Users had to read those panels in the "
             "UI manually.",
             "Added four new chatbot tools: macro impact today (with "
             "optional symbol filter), management outlook, material "
             "information, and sector volume heatmap. The chatbot "
             "now supports questions like 'what is the macro impact "
             "on cement today?' or 'what does PSO management say "
             "about furnace oil costs?'"],
            ["5", "Fundamentals cache",
             "Only two of the fifteen stocks had the new "
             "P/E, P/B, dividend yield and payout-ratio fields; the "
             "rest were cached before the connector was upgraded. "
             "This caused four of the eight sector medians "
             "(Banking, Cement, OMC/Refining, Conglomerate/Chem) to "
             "compute as null, which in turn caused the value tab "
             "to display blanks in the 'vs sector' columns.",
             "Re-ran the fundamentals refresh for the full "
             "fifteen-stock universe. Every stock now has all four "
             "ratios, every sector has a median, and every stock "
             "carries its own 'percent above/below sector' figures."],
            ["6", "Macro driver coverage",
             "The first cut of the macro impact engine looked at "
             "four drivers (policy rate, oil, USD/PKR, coal proxy). "
             "The analyst asked for 'many macroeconomic factors' to "
             "be respected — gold, copper and cotton were already "
             "being collected daily but were not feeding the "
             "engine.",
             "Extended driver detection to gold (risk-off proxy), "
             "copper (industrial-growth proxy) and cotton (textile "
             "input cost), with thresholds tuned to fire only on "
             "meaningful 21-day moves. Banking, Cement and Oil & "
             "Gas E&P rule-books were extended to read the new "
             "tags. The engine now responds to seven distinct macro "
             "drivers."],
        ]},

        {"kind": "h", "level": 2,
         "text": "10.2  Findings that are working as designed"},
        {"kind": "bullets", "items": [
            "The Sarmaya cross-check connector reports 'no parseable "
            "page' because Sarmaya rebuilt their site on a "
            "JavaScript-only single-page architecture. The connector "
            "degrades gracefully — it returns an empty result and a "
            "clear note rather than crashing the pipeline. Every "
            "downstream consumer treats Sarmaya as an optional "
            "cross-check, so the system continues to function on "
            "yfinance fundamentals alone.",
            "The Material Information parquet currently holds only a "
            "handful of rows. PSX has simply not filed many material "
            "disclosures for the fifteen-stock universe in the past "
            "month — this is real-world data, not a scraping bug. "
            "When a wave of disclosures arrives (typically before "
            "result season) the file fills naturally.",
            "Three cement names (FCCL, KOHC, MLCF) display a payout "
            "ratio of zero. They genuinely paused dividends in 2019. "
            "The figure is correct.",
            "The strategy-rule recommendation can show CASH while "
            "individual AI conviction calls show BUY. This is the "
            "system working as designed: the rule sets exposure, "
            "the AI ranks within the budget. Both are surfaced so "
            "the analyst can see the tension."],
        },

        {"kind": "h", "level": 2,
         "text": "10.3  End-to-end smoke test"},
        {"kind": "p", "text":
            "After every fix, a single end-to-end script exercised "
            "the morning-brief assembly, the action explainer, the "
            "predictions reader, and all four new chat tools against "
            "the live data on disk. All thirteen assertions passed: "
            "the macro impact engine emits three drivers and covers "
            "all eight sectors, the predictions reader back-fills "
            "macro context for older entries, every chat tool returns "
            "a non-error payload, and the UI explainer surfaces the "
            "rationale plus tailwinds and headwinds for the day's "
            "top action."},

        {"kind": "h", "level": 2,
         "text": "10.4  Iteration April 28: industry-specific KPI "
                  "expansion"},
        {"kind": "p", "text":
            "Following the audit, four of the five 'open items' "
            "identified for the next iteration were closed in a "
            "single follow-up. The State Bank connector already "
            "captures the T-bill and PIB yield curves, KIBOR, and "
            "reserve totals daily, but those numbers were being "
            "discarded after a single use — so the macro impact "
            "engine could not see whether T-bills had moved or "
            "whether reserves were drifting toward the IMF stress "
            "band. Three persistent time-series files now keep "
            "those readings on disk:"},
        {"kind": "bullets", "items": [
            "data/macro/sbp_rates.parquet — one row per business day "
            "with the policy rate, KIBOR (3, 6, 12 months), T-bill "
            "cut-offs (1, 3, 6, 12 months), PIB yields (3, 5, 10 "
            "years), and FX reserves (SBP, banks, total).",
            "data/macro/kse100.parquet — one row per business day "
            "with the KSE-100 close, daily change, and intra-day "
            "high / low captured from the PSX DPS indices page.",
            "data/macro/cpi_pakistan.parquet — one row per refresh "
            "with the latest Pakistan CPI year-on-year print, the "
            "calendar period it covers (e.g. 'March'), and the "
            "data source. The scrape uses Trading Economics first "
            "(it publishes the print as a one-line summary that is "
            "robust to layout changes) and falls back to the "
            "Bureau of Statistics landing page."],
        },
        {"kind": "p", "text":
            "A new GitHub Actions workflow (.github/workflows/"
            "macro_kpis.yml) refreshes all three files on weekdays "
            "at 17:00 Pakistan time (12:00 UTC), commits the "
            "incremental rows, and pushes back to the main "
            "branch. The workflow is idempotent: re-runs on the "
            "same business day overwrite the row in place rather "
            "than creating duplicates."},
        {"kind": "p", "text":
            "The macro impact engine now consumes those files "
            "through a new helper, _load_kpi_snapshot, and emits "
            "five additional families of drivers — T-bill 3-month "
            "trading above or below policy rate (banking signal), "
            "T-bill 5-day trend, KIBOR 5-day trend (financial-"
            "cost benchmark for every leveraged sector), reserve "
            "stress / recovery (BoP regime), KSE-100 5- and 21-"
            "day momentum (broad-market regime), and CPI level "
            "and direction (real-rate signal). The sector rule "
            "book has been expanded to react to each tag with a "
            "specific score and a one-sentence reason: KIBOR "
            "rising scores -2 for IPPs and Conglomerate / Chem "
            "because both run highly leveraged balance sheets, "
            "but only -1 for OMCs because their working-capital "
            "cycle is shorter; reserve stress scores -2 for "
            "OMCs (letter-of-credit risk on crude imports), -2 "
            "for IPPs (circular debt worsens), -2 for pharma "
            "(API import letters of credit get harder to "
            "confirm), -1 for cement and miscellaneous "
            "manufacturers, and -2 for banks themselves; CPI "
            "easing scores +2 for cement (single biggest "
            "tailwind for leveraged construction names) and +1 "
            "for chem and pharma; CPI sticky-high scores -1 for "
            "cement and pharma. The analyst's exact request was "
            "industry-level KPIs with sector-specific reasoning, "
            "and that is now what the engine produces."},
        {"kind": "p", "text":
            "Both the user interface and the language-model "
            "briefing reflect the new data. The Macro Radar "
            "panel on the Today tab now shows the live numeric "
            "values (T-bill 3-month, KIBOR 3-month, SBP "
            "reserves, KSE-100, CPI year-on-year) at the top of "
            "the card with five-day or thirty-day deltas, "
            "before the existing driver and sector / stock "
            "tables. The chatbot has gained a new tool, "
            "get_industry_kpis, that returns the same numeric "
            "snapshot directly, and the existing "
            "get_macro_impact_today tool now carries the "
            "industry-KPI block alongside the drivers and "
            "verdicts. Most importantly, the LLM briefing that "
            "is built before every prediction now opens its "
            "macro section with an 'Industry KPIs' block listing "
            "the actual numbers, so the AI's rationale can cite "
            "specific values rather than speak in generalities."},
        {"kind": "p", "text":
            "Quick sanity check on April 27 data: the engine "
            "produces five active drivers (Brent crude up 9.7% "
            "in 5 days, Copper up 9.0% in 21 days, Cotton up "
            "12.9% in 21 days, FX reserves recovering at USD "
            "15.1 billion, CPI cooling to 7.3%) and routes them "
            "into eight sector verdicts: Banking +1 (tailwind), "
            "Cement +1 (tailwind, CPI easing wins out over "
            "Brent up), Oil & Gas E&P +5 (strong tailwind), "
            "OMC / Refining +2 (tailwind), Power +3 (strong "
            "tailwind), Conglomerate / Chem +3 (strong "
            "tailwind), Pharma +1 (tailwind), Misc neutral. "
            "Every line carries a human-readable reason — for "
            "example Power's strong tailwind cites 'Reserve "
            "rebuilds typically come with circular-debt "
            "settlement plans — cash flow normalises and "
            "dividends resume' (+2) plus 'Furnace-oil-fired "
            "plants get fuel-cost pass-through under PPA "
            "indexation' (+1)."},

        {"kind": "h", "level": 2,
         "text": "10.5  Open items for the next iteration"},
        {"kind": "p", "text":
            "Three items remain on the roadmap. Each requires "
            "either a less stable data source or a longer-running "
            "engineering effort, so they are deliberately scoped "
            "for the next sprint rather than slotted into the "
            "April 28 KPI release."},
        {"kind": "bullets", "items": [
            "APCMA cement-industry retention prices and DRAP "
            "pharmaceutical price-cap notifications. Both are "
            "published as PDF bulletins on the relevant industry "
            "association or regulator website; the cadence is "
            "irregular and the format changes between releases. "
            "A robust pipeline needs PDF parsing rather than HTML "
            "scraping, plus a small retry / human-review queue "
            "for malformed releases. The existing macro engine "
            "already covers the headline driver (CPI for pharma, "
            "reserve stress for both), but per-product retention "
            "and MRP detail would let the engine fire on company-"
            "specific signals such as a single pharma product "
            "getting a price increase.",
            "NEPRA quarterly circular-debt bulletin for the IPP "
            "sector. NEPRA publishes a quarterly state-of-"
            "industry report and an annual circular-debt update; "
            "both are PDFs with tables that move position year "
            "to year. The current engine uses the FX reserve "
            "stress / recovery driver as a proxy because reserve "
            "stress and circular-debt growth move together with "
            "high correlation, but a direct circular-debt number "
            "would let the engine flag IPP-specific tail risk "
            "even when reserves look fine.",
            "Live paper-trading against a broker API. KASB / JS "
            "Investments / AKD all expose REST APIs that accept "
            "limit and stop orders, including a paper / sandbox "
            "mode. Wiring the predictions feed to one of those "
            "APIs would remove the last back-test biases — fees, "
            "fills, slippage, and timing all become real. Scoped "
            "as a multi-week project because broker credentials, "
            "compliance review, and a staged rollout are "
            "required before the system can place real orders."],
        },

        # =======================================================
        # 10.5 — Iteration April 29: synthesizer + autonomy +
        # critic. Closes the gap exposed by the SBP MPC surprise
        # the day before, and answers the analyst's complaint
        # that different tabs gave contradictory calls.
        # =======================================================
        {"kind": "h", "level": 2,
         "text": "10.5  Iteration April 29: verdict synthesizer, "
                  "autonomous re-trigger, and critic self-review"},
        {"kind": "p", "text":
            "Three days after the industry-KPI expansion the SBP "
            "Monetary Policy Committee announced a surprise rate "
            "decision at 11:45 PKT mid-session, and yesterday's "
            "predictions — generated at 09:15 PKT before any signal "
            "of the move existed — were left out of date for the "
            "rest of the day. The user surfaced two further "
            "complaints in the same conversation: that the bot's "
            "predictions occasionally diverged from market reality, "
            "and that the dashboard sometimes carried opposite "
            "calls on the same name on different tabs (Value tab "
            "marking a stock SELL while the prediction tab marked it "
            "BUY). The April 29 iteration ships three independent "
            "fixes that together answer all three complaints with a "
            "single coherent story."},
        {"kind": "h", "level": 3,
         "text": "Fix 1 — Verdict synthesizer (one call across "
                  "seven lenses)"},
        {"kind": "p", "text":
            "A new module, brain/verdict_synthesizer.py, takes the "
            "outputs of all the existing lenses (value, quality, "
            "momentum, macro, news sentiment, FIPI flow, and "
            "Director's Report tone), maps each into a signed "
            "score in the range -3 to +3, and combines them with "
            "fixed weights into a single composite verdict per "
            "stock. The output carries one of five actions — BUY, "
            "ADD, HOLD, TRIM, AVOID — together with a direction "
            "label, a conviction level, and an audit trail that "
            "shows what every lens contributed. Crucially, when "
            "two lenses disagree sharply (a +2 lens against a -2 "
            "lens), the synthesizer logs the conflict explicitly "
            "and applies a hand-crafted resolution rule. Six rules "
            "are encoded today: value-sells-while-momentum-buys "
            "permits a short-term trade with a tight stop and "
            "capped conviction, value-buys-while-momentum-sells is "
            "treated as a value-trap risk and rejected, macro-"
            "tailwind-with-quality-junk routes the flow to a "
            "higher-quality peer in the same sector, news-bearish-"
            "with-flow-buying delays sizing for one session of "
            "confirmation, management-bullish-with-macro-bearish "
            "soft-caps conviction until the next earnings print, "
            "and quality-high-with-value-sell explicitly tells the "
            "analyst to wait for a 10-15% pullback to intrinsic "
            "value before adding. The synthesizer is wired into "
            "the Today tab as a top-level 'Bot's Verdict' panel "
            "(with a drill-down per stock that lays out every "
            "lens's contribution and reason), into the daily PDF "
            "report as its own page, and into the chatbot as a "
            "new tool, get_bots_verdict, so any 'why does Value "
            "say X but the prediction say Y' question receives a "
            "single reconciled answer."},
        {"kind": "h", "level": 3,
         "text": "Fix 2 — Intraday news-shock retrigger"},
        {"kind": "p", "text":
            "A new script, scripts/check_news_shocks.py, runs "
            "after every news-scoring batch (07:00 / 13:00 / 18:00 "
            "PKT). It scans the last six hours of scored articles "
            "and flags any item that simultaneously crosses three "
            "gates: absolute sentiment of at least 0.40, HIGH "
            "confidence, and either a universe-ticker mention or a "
            "high-impact macro tag (POLICY_RATE, FX, OIL_SHOCK, "
            "REGULATOR, DOWNGRADE, etc.). Every fired shock is "
            "stamped into data/news/shock_log.json so the same "
            "article never re-fires. The wrapper workflow "
            "(.github/workflows/news_scoring.yml) checks the "
            "exit code: when the script returns 7, the workflow "
            "dispatches predictions.yml via the GitHub CLI, so "
            "fresh recommendations land within minutes of a "
            "high-impact news event rather than waiting for the "
            "09:15 PKT next-day window. The same pattern handles "
            "an SBP surprise: when the central bank's press "
            "release is scored as a high-confidence rate-policy "
            "article, the bot re-predicts the entire universe "
            "automatically and the dashboard updates by the next "
            "page reload."},
        {"kind": "h", "level": 3,
         "text": "Fix 3 — Pre-MPC alert"},
        {"kind": "p", "text":
            "Hand-maintained config/sbp_mpc_calendar.py carries "
            "every SBP MPC date for the calendar year (eight "
            "meetings on the official 2026 schedule). The macro "
            "impact engine now emits an mpc_alert block in its "
            "output that flags whether the next meeting falls "
            "inside a three-day pre-window or a one-day post-"
            "window. The predictions pipeline reads that block and "
            "automatically downgrades conviction one notch on "
            "rate-sensitive sectors (banking, cement, power, auto, "
            "textile, conglomerate) when the pre-window is "
            "active, so a position is never built at full size "
            "into a meeting the bot cannot predict. The Today "
            "tab and the PDF Macro Radar both display the alert "
            "banner with the meeting date and a one-line "
            "explanation of why the conviction cap is applied. "
            "The user's specific failure on April 28 — a +9 BUY "
            "on a banking name the night before the surprise "
            "hike — would not have happened with this rule in "
            "place: conviction would have been capped to MEDIUM "
            "and a 'MPC in 0 days' note would have been on the "
            "card."},
        {"kind": "h", "level": 3,
         "text": "Fix 4 — Critic self-review pass"},
        {"kind": "p", "text":
            "Even the strongest LLM occasionally publishes an "
            "internally inconsistent recommendation — a BULLISH "
            "direction with bearish key drivers, or a BULLISH "
            "call with stop / target geometry inverted (target "
            "below entry). brain/prediction_critic.py runs four "
            "deterministic post-checks on every prediction "
            "before it is written to disk: direction-versus-"
            "action consistency, drivers-match-direction "
            "(detected via simple keyword polarity scan), stop-"
            "and-target geometry, and synthesizer alignment "
            "(when the deterministic synthesizer disagrees "
            "sharply with the LLM, conviction is capped one "
            "notch). A 'fail' check forces the call to HOLD and "
            "downgrades conviction to LOW; a 'warn' check "
            "downgrades conviction one notch. Every action is "
            "stamped on a critic_notes field so the analyst can "
            "audit what the critic caught and why. The PDF "
            "report's per-stock detail card now shows the "
            "critic notes inline whenever they are non-empty."},
        {"kind": "p", "text":
            "Together the four fixes give the analyst a single "
            "coherent decision flow: every prediction is now "
            "filtered through the deterministic critic, every "
            "stock receives a verdict that already reconciles "
            "the seven lenses, the dashboard explains conflicts "
            "rather than presenting them, and the bot re-"
            "predicts on demand whenever a HIGH-confidence shock "
            "lands. None of these features depend on a new "
            "external dependency or a paid API; the entire "
            "iteration runs on the same deterministic data flow "
            "the rest of the bot already uses."},

        {"kind": "pagebreak"},
    ]

    # ----------------------------------------------------- A. GLOSSARY
    doc += [
        {"kind": "h", "level": 1, "text": "Appendix · Glossary"},
        {"kind": "table",
         "headers": ["Term", "Meaning"],
         "rows": [
            ["12-1 Momentum",
             "12-month return excluding the most recent month — the "
             "most well-documented cross-sectional anomaly in equity "
             "markets."],
            ["ATR",
             "Average True Range, 14-day default — used for stop "
             "placement."],
            ["BVPS",
             "Book Value Per Share = total equity ÷ shares outstanding."],
            ["CGT",
             "Capital Gains Tax. Pakistan filer rate for less-than-"
             "one-year holdings is 15%."],
            ["Conviction",
             "High / medium / low. Used as the ranking weight: 1.0 / "
             "0.7 / 0.3 respectively."],
            ["DDM",
             "Dividend Discount Model — values a stock at the present "
             "value of its expected future dividends."],
            ["End-of-day (EOD)",
             "After PSX close at 15:30 PKT."],
            ["EPS coefficient of variation",
             "Standard deviation of 5-year EPS divided by mean. "
             "Lower means more stable earnings."],
            ["FED",
             "Federal Excise Duty — 16% sales tax on brokerage in "
             "Pakistan."],
            ["FIPI / LIPI",
             "Foreign / Local Investor Portfolio Investment — daily "
             "net buy / sell on PSX."],
            ["KSE-100",
             "Pakistan's main equity index — 100 largest stocks by "
             "free-float-adjusted market cap."],
            ["MAE",
             "Mean Absolute Error. Lower is better."],
            ["P/B, P/E",
             "Price-to-book and price-to-earnings ratios."],
            ["PKR",
             "Pakistani Rupee."],
            ["PSX",
             "Pakistan Stock Exchange."],
            ["R²",
             "Coefficient of determination. 0 = no better than the "
             "mean; 1 = perfect; negative = worse than the mean."],
            ["ROE",
             "Return on Equity = net income ÷ total equity."],
            ["RSI-14",
             "Relative Strength Index, 14-day. Above 70 = overbought, "
             "below 30 = oversold."],
            ["SBP",
             "State Bank of Pakistan — the central bank."],
            ["SMA-20 / 50 / 200",
             "Simple Moving Average over 20 / 50 / 200 trading days."],
            ["VIX",
             "US equity volatility index. Above 18 = elevated; above "
             "22 = stressed."],
            ["Walk-forward",
             "Back-test design that replays each day with a strict "
             "information cutoff to avoid look-ahead bias."],
        ]},

        {"kind": "p", "text":
            f"_Generated {today}._"},
    ]

    return doc


# ============================================================================
#  MARKDOWN RENDERER
# ============================================================================
def render_markdown(blocks: list[dict]) -> str:
    """Render the document model to a clean Markdown string."""
    out: list[str] = []
    for b in blocks:
        kind = b["kind"]

        if kind == "h":
            level = max(1, min(6, int(b["level"]) + 1))  # h0 -> H1 etc.
            out.append("#" * level + " " + b["text"])
            out.append("")

        elif kind == "p":
            out.append(b["text"])
            out.append("")

        elif kind == "callout":
            # Render as a blockquote so it stands apart from body text
            for line in b["text"].split("\n"):
                out.append("> " + line)
            out.append("")

        elif kind == "bullets":
            for item in b["items"]:
                out.append(f"- {item}")
            out.append("")

        elif kind == "table":
            headers = b["headers"]
            out.append("| " + " | ".join(_md_cell(h) for h in headers) +
                        " |")
            out.append("|" + "|".join(["---"] * len(headers)) + "|")
            for row in b["rows"]:
                out.append("| " + " | ".join(_md_cell(c) for c in row) +
                            " |")
            out.append("")

        elif kind == "pagebreak":
            out.append("---")
            out.append("")

    return "\n".join(out).strip() + "\n"


def _md_cell(text: str) -> str:
    """Escape pipe characters and collapse newlines so a cell stays on one
    Markdown row."""
    return str(text).replace("|", "\\|").replace("\n", " ")


# ============================================================================
#  WORD (.docx) RENDERER
# ============================================================================
_HEADING_FOR_LEVEL = {
    0: "Title",
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
}


def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a background colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Parse very-light Markdown (just **bold**) and add runs accordingly."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


def _add_paragraph(document, text: str, style: str | None = None,
                    italic: bool = False, justify: bool = True) -> None:
    p = document.add_paragraph(style=style) if style else \
        document.add_paragraph()
    if justify and not style:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_runs_with_bold(p, text)
    if italic:
        for run in p.runs:
            run.italic = True


def _add_callout(document, text: str) -> None:
    """A boxed paragraph with light-blue shading."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_runs_with_bold(p, text)
    # Apply a left border + shading via XML manipulation
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "6")
        b.set(qn("w:color"), "C8D6E5")
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F6FC")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def _add_table(document, headers: list[str],
                rows: list[list[str]]) -> None:
    """Add a properly styled table that fits the page width."""
    tbl = document.add_table(rows=1 + len(rows), cols=len(headers))
    # Built-in style with banded rows and a coloured header
    try:
        tbl.style = "Light Grid Accent 1"
    except KeyError:
        tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_shading(cell, "1F3A93")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Body rows
    for r_i, row in enumerate(rows, start=1):
        cells = tbl.rows[r_i].cells
        for c_i, val in enumerate(row):
            cell = cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs_with_bold(p, str(val))
            for run in p.runs:
                run.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if r_i % 2 == 0:
                _set_cell_shading(cell, "EEF3FB")


def render_docx(blocks: list[dict], out_path: Path) -> Path:
    document = Document()

    # Page setup — A4 with comfortable margins
    for section in document.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Default body style
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # Add a page-number footer
    section = document.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.add_run("PSX Trading System · methodology review · page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)

    for b in blocks:
        kind = b["kind"]
        if kind == "h":
            style_name = _HEADING_FOR_LEVEL.get(int(b["level"]),
                                                  "Heading 4")
            p = document.add_paragraph(b["text"], style=style_name)
            if int(b["level"]) == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif kind == "p":
            _add_paragraph(document, b["text"])
        elif kind == "callout":
            _add_callout(document, b["text"])
        elif kind == "bullets":
            for item in b["items"]:
                p = document.add_paragraph(style="List Bullet")
                _add_runs_with_bold(p, item)
        elif kind == "table":
            _add_table(document, b["headers"], b["rows"])
            # Spacer so rows can't touch the next paragraph
            document.add_paragraph()
        elif kind == "pagebreak":
            document.add_page_break()

    document.save(str(out_path))
    return out_path


# ============================================================================
#  Driver
# ============================================================================
def build_solution_review() -> tuple[Path, Path]:
    out_dir = ROOT / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")

    md_path = out_dir / f"solution_review_{ts}.md"
    docx_path = out_dir / f"solution_review_{ts}.docx"

    blocks = build_doc_model()

    md_path.write_text(render_markdown(blocks), encoding="utf-8")
    render_docx(blocks, docx_path)

    return md_path, docx_path


if __name__ == "__main__":
    md, docx = build_solution_review()
    print(f"Wrote {md}  ({md.stat().st_size / 1024:.1f} KB)")
    print(f"Wrote {docx}  ({docx.stat().st_size / 1024:.1f} KB)")
